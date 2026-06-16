"""
DOCX document processing for Nexora001. 
Extracts text from Word documents using python-docx. 
"""

from docx import Document
from pathlib import Path
from typing import Dict, Optional
import sys

sys.path.insert(0, str(Path(__file__).parent.parent.parent))

from nexora001.processors. chunker import TextChunker
from nexora001.storage.mongodb import get_storage
from nexora001.processors.embeddings import get_embedding_generator


class DOCXProcessor:
    """Process DOCX documents and store in vector database."""
    
    def __init__(
        self,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        generate_embeddings: bool = True
    ):
        """
        Initialize DOCX processor.
        
        Args:
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between chunks
            generate_embeddings: Whether to generate embeddings
        """
        self.chunker = TextChunker(chunk_size, chunk_overlap)
        self.generate_embeddings = generate_embeddings
        
        if generate_embeddings:
            self.embedding_generator = get_embedding_generator("sentence_transformers")
    
    def extract_text(self, docx_path: str) -> Dict:
        """
        Extract text from a DOCX file.
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            Dictionary with text and metadata
        """
        docx_path = Path(docx_path)
        
        if not docx_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {docx_path}")
        
        try:
            # Open document
            doc = Document(str(docx_path))
            
            # Extract text from paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text. strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            full_text = "\n\n".join(text_parts)
            
            # Extract metadata
            core_properties = doc.core_properties
            metadata = {
                "filename": docx_path.name,
                "title": core_properties.title or docx_path.stem,
                "author": core_properties.author or "Unknown",
                "subject": core_properties.subject or "",
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables)
            }
            
            return {
                "text": full_text,
                "metadata": metadata
            }
            
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX: {e}")
    
    def process_and_store(self, docx_path: str, client_id: str, source_url: Optional[str] = None) -> Dict:
        """
        Process DOCX and store in database. 
        
        Args:
            docx_path: Path to DOCX file
            client_id: Client ID for multi-tenant data isolation
            source_url: Optional URL/identifier for the document
            
        Returns:
            Dictionary with processing statistics
        """
        # Extract text
        result = self.extract_text(docx_path)
        text = result["text"]
        metadata = result["metadata"]
        
        if not text or len(text.strip()) < 100:
            return {
                "success": False,
                "error": "Insufficient text content in DOCX",
                "chunks_created": 0
            }
        
        # Determine source URL
        if not source_url:
            source_url = f"file://{Path(docx_path). absolute()}"
        
        # Chunk the text
        chunks = self.chunker.chunk_text(text, metadata={
            "source_url": source_url,
            "title": metadata["title"]
        })
        
        # Store chunks
        with get_storage() as storage:
            chunks_stored = 0
            
            for chunk in chunks:
                try:
                    chunk_text = chunk['text']
                    chunk_index = chunk['chunk_index']
                    total_chunks = chunk['total_chunks']
                    
                    # Generate embedding if enabled
                    embedding = None
                    if self.generate_embeddings:
                        embedding = self.embedding_generator.generate_embedding(chunk_text)
                    
                    # Store in MongoDB
                    if embedding:
                        storage.store_document_with_embedding(
                            client_id=client_id,
                            content=chunk_text,
                            embedding=embedding,
                            source_url=source_url,
                            source_type="docx",
                            title=metadata["title"],
                            chunk_index=chunk_index,
                            total_chunks=total_chunks,
                            metadata={
                                "filename": metadata["filename"],
                                "author": metadata["author"],
                                "paragraphs": metadata["paragraphs"],
                                "tables": metadata["tables"],
                                "chunk_char_count": chunk['char_count']
                            }
                        )
                    else:
                        storage.store_document(
                            client_id=client_id,
                            content=chunk_text,
                            source_url=source_url,
                            source_type="docx",
                            title=metadata["title"],
                            metadata={
                                "filename": metadata["filename"],
                                "author": metadata["author"],
                                "chunk_index": chunk_index,
                                "total_chunks": total_chunks,
                                "chunk_char_count": chunk['char_count']
                            }
                        )
                    
                    chunks_stored += 1
                    
                except Exception as e:
                    print(f"Failed to store chunk {chunk_index}: {e}")
                    continue
        
        return {
            "success": True,
            "filename": metadata["filename"],
            "title": metadata["title"],
            "paragraphs": metadata["paragraphs"],
            "chunks_created": chunks_stored,
            "total_characters": len(text)
        }


def process_docx(docx_path: str, source_url: Optional[str] = None, client_id: Optional[str] = None) -> Dict:
    """
    Convenience function to process a DOCX file.
    
    Args:
        docx_path: Path to DOCX file
        source_url: Optional URL/identifier
        client_id: Client ID for multi-tenant support
        
    Returns:
        Processing statistics
    """
    processor = DOCXProcessor()
    return processor.process_and_store(docx_path, client_id, source_url)